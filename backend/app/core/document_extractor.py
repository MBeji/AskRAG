"""
Service d'extraction de documents multi-format
Étape 14.5: Pipeline d'extraction de contenu (PDF, DOCX, TXT, MD, HTML)
"""

import os
import io
import logging
import mimetypes
from typing import Dict, Any, Optional, List, Union
from pathlib import Path

# Imports pour extraction
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False


class DocumentExtractor:
    """
    Service d'extraction de contenu à partir de documents multi-format
    Support PDF, DOCX, TXT, MD, HTML
    """
    
    def __init__(self):
        """Initialise le service d'extraction"""
        self.logger = logging.getLogger(__name__)
        
        # Types de fichiers supportés
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,  # Tentative avec docx
            '.txt': self._extract_text,
            '.md': self._extract_markdown,
            '.markdown': self._extract_markdown,
            '.html': self._extract_html,
            '.htm': self._extract_html
        }
        
        self.logger.info("DocumentExtractor initialisé")
    
    def detect_format(self, file_path: Union[str, Path] = None, 
                     file_content: bytes = None,
                     filename: str = None) -> Optional[str]:
        """
        Détecte le format d'un fichier
        
        Args:
            file_path: Chemin du fichier
            file_content: Contenu binaire du fichier
            filename: Nom du fichier pour l'extension
            
        Returns:
            str: Extension détectée ou None
        """
        try:
            # Détecter par extension de fichier
            if file_path:
                return Path(file_path).suffix.lower()
            elif filename:
                return Path(filename).suffix.lower()
            
            # Détecter par type MIME si contenu disponible
            if file_content and MAGIC_AVAILABLE:
                mime_type = magic.from_buffer(file_content, mime=True)
                
                mime_to_ext = {
                    'application/pdf': '.pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                    'application/msword': '.doc',
                    'text/plain': '.txt',
                    'text/markdown': '.md',
                    'text/html': '.html'
                }
                
                return mime_to_ext.get(mime_type)
            
            # Fallback sur mimetypes
            if filename:
                mime_type, _ = mimetypes.guess_type(filename)
                if mime_type == 'application/pdf':
                    return '.pdf'
                elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    return '.docx'
                elif mime_type == 'text/plain':
                    return '.txt'
                elif mime_type == 'text/html':
                    return '.html'
            
            return None
            
        except Exception as e:
            self.logger.error(f"Erreur détection format: {e}")
            return None
    
    def extract_content(self, 
                       file_path: Union[str, Path] = None,
                       file_content: bytes = None,
                       filename: str = None,
                       format_hint: str = None) -> Dict[str, Any]:
        """
        Extrait le contenu d'un document
        
        Args:
            file_path: Chemin du fichier
            file_content: Contenu binaire du fichier
            filename: Nom du fichier
            format_hint: Format suggéré (extension)
            
        Returns:
            Dict: Contenu extrait et métadonnées
        """
        try:
            # Détecter le format
            file_format = format_hint or self.detect_format(
                file_path=file_path,
                file_content=file_content,
                filename=filename
            )
            
            if not file_format:
                raise ValueError("Format de fichier non détecté")
            
            if file_format not in self.supported_formats:
                raise ValueError(f"Format non supporté: {file_format}")
            
            # Préparer les données
            if file_path:
                with open(file_path, 'rb') as f:
                    content_bytes = f.read()
                filename = filename or Path(file_path).name
            elif file_content:
                content_bytes = file_content
                filename = filename or "document"
            else:
                raise ValueError("Aucun contenu fourni")
            
            # Extraire selon le format
            extractor = self.supported_formats[file_format]
            content = extractor(content_bytes)
            
            # Métadonnées de base
            metadata = {
                'filename': filename,
                'format': file_format,
                'size_bytes': len(content_bytes),
                'content_length': len(content),
                'extracted_at': None,  # Sera ajouté par le service appelant
                'extraction_success': True
            }
            
            self.logger.info(f"Contenu extrait: {filename} ({file_format})")
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except Exception as e:
            self.logger.error(f"Erreur extraction: {e}")
            return {
                'content': '',
                'metadata': {
                    'filename': filename or 'unknown',
                    'format': file_format or 'unknown',
                    'extraction_success': False,
                    'error': str(e)
                }
            }
    
    def _extract_pdf(self, content: bytes) -> str:
        """Extrait le texte d'un PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Erreur extraction PDF: {e}")
            raise
    
    def _extract_docx(self, content: bytes) -> str:
        """Extrait le texte d'un document DOCX/DOC"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extraire aussi les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Erreur extraction DOCX: {e}")
            raise
    
    def _extract_text(self, content: bytes) -> str:
        """Extrait le texte d'un fichier texte"""
        try:
            # Essayer différents encodages
            for encoding in ['utf-8', 'utf-16', 'latin1', 'cp1252']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # Fallback avec erreurs ignorées
            return content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            self.logger.error(f"Erreur extraction texte: {e}")
            raise
    
    def _extract_markdown(self, content: bytes) -> str:
        """Extrait le texte d'un fichier Markdown"""
        try:
            # Décoder le contenu
            md_text = self._extract_text(content)
            
            # Convertir en HTML puis extraire le texte
            html = markdown.markdown(md_text, extensions=['tables', 'codehilite'])
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extraire le texte en préservant la structure
            text_parts = []
            for element in soup.descendants:
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    text_parts.append(f"\n## {element.get_text().strip()}\n")
                elif element.name == 'p':
                    text = element.get_text().strip()
                    if text:
                        text_parts.append(text)
                elif element.name == 'li':
                    text = element.get_text().strip()
                    if text:
                        text_parts.append(f"- {text}")
                elif element.name == 'code':
                    text = element.get_text().strip()
                    if text:
                        text_parts.append(f"`{text}`")
            
            return '\n\n'.join(text_parts) if text_parts else md_text
            
        except Exception as e:
            self.logger.error(f"Erreur extraction Markdown: {e}")
            # Fallback sur texte brut
            return self._extract_text(content)
    
    def _extract_html(self, content: bytes) -> str:
        """Extrait le texte d'un fichier HTML"""
        try:
            # Décoder le contenu
            html_text = self._extract_text(content)
            
            # Parser avec BeautifulSoup
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Supprimer les scripts et styles
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extraire le texte principal
            text_parts = []
            
            # Titres
            for header in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                text = header.get_text().strip()
                if text:
                    text_parts.append(f"## {text}")
            
            # Paragraphes
            for para in soup.find_all('p'):
                text = para.get_text().strip()
                if text:
                    text_parts.append(text)
            
            # Listes
            for li in soup.find_all('li'):
                text = li.get_text().strip()
                if text:
                    text_parts.append(f"- {text}")
            
            # Tables
            for table in soup.find_all('table'):
                for row in table.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if any(cells):
                        text_parts.append(' | '.join(cells))
            
            # Si aucun contenu structuré, prendre tout le texte
            if not text_parts:
                text_parts = [soup.get_text().strip()]
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Erreur extraction HTML: {e}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Retourne la liste des formats supportés"""
        return list(self.supported_formats.keys())
    
    def is_supported(self, format_or_filename: str) -> bool:
        """Vérifie si un format/fichier est supporté"""
        if format_or_filename.startswith('.'):
            return format_or_filename.lower() in self.supported_formats
        else:
            extension = Path(format_or_filename).suffix.lower()
            return extension in self.supported_formats


# Instance globale
document_extractor = DocumentExtractor()
